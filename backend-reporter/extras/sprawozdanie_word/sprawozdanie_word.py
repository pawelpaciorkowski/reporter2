import os.path
from dataclasses import dataclass
from typing import Any, Dict

import code128
from docx import Document
from docx.blkcntnr import BlockItemContainer
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.shared import Mm, Pt, RGBColor
from docx.table import Table

from datasources.centrum import Centrum
from datasources.spreadsheet import spreadsheet_to_values
from helpers import empty, copy_from_remote
from helpers.crystal_ball.marcel_servers import sciezka_wydruku
from helpers.marcel_pdf_xml import MarcelPdf
from extras.sprawozdanie_word.docx_mods import add_page_number

from extras.sprawozdanie_word import SprawozdanieZCentrum

"""
sposób wydruku badania : 1 nie drukować, 2 drukować odstęp, 3 drukować nazwę

TODO

[09:28] Aleksandra Wilczak
marginesy - w wordzie, zeby bylo miejsce na stopkę z numerami stron 1/4 , 2/4 itd.
[09:39] Aleksandra Wilczak
brakuje materiału 
[09:40] Aleksandra Wilczak
i daty godz pobrania


Kody przykładów

6086028260
6086028270
6098716280




"""

STOPKA = """Without the written consent of the Laboratory, the test report may not be reproduced except in full. The \
test results refer only to sample whose type, date and time of collection, date and time of acceptance for \
testing are identified in this report. People who have a patient account at ALAB laboratories can benefit \
from specialist support in the field of preventive consultation and interpretation result as part of the "Ask \
the laboratory" service. \
Patient identification is based on the first 9 digits of the barcode (counting from the left). The 10th digit is \
the digit intended for lab. More information at: https://sklep.alablaboratoria.pl/922-przygotowanie-do-badania."""

BADANIE_KLINICZNE = """\nThe examination was performed for the purposes of clinical trials."""

PODPIS = """\nThis printout is information about the laboratory test report. The report was
prepared in electronic form signed with a qualified electronic signature.
Authorized by: Laboratory Diagnostician ............ PWZDL No. ............\n"""


def laduj_tlumaczenia(slownik_fn):
    res = {'sek': {}, 'bad': {}, 'mat': {}, 'par': {}, 'nor': {}, 'wyn': {}}
    for row in spreadsheet_to_values(slownik_fn)[1:]:
        for k in res.keys():
            if row[0] is not None and row[0].lower().startswith(k):
                if row[1] is not None and row[2] is not None:
                    translate_from = row[1].strip()
                    if k not in ('nor',):
                        translate_from = translate_from.upper()
                    res[k][translate_from] = row[2].strip()
    return res


@dataclass
class SprawozdanieWordConfig:
    kliniczne: bool = False


class SprawozdanieWord:
    def __init__(self, szc: SprawozdanieZCentrum, slownik_fn, config: SprawozdanieWordConfig = None):
        self.config: SprawozdanieWordConfig = config
        self.szc = szc
        self.sekcja = None
        self.probka = None
        self.wykonanie = None
        self.section_header_printed = False
        self.files_to_clean_up = []
        self._tlumaczenia = laduj_tlumaczenia(slownik_fn)

    def _setup_styles(self):
        styles = self.doc.styles
        styles['Heading 1'].alignment = WD_ALIGN_PARAGRAPH.CENTER

        style = styles.add_style('OpisMetody', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']

    def _create_compact_table(self, rows: int, cols: int) -> Table:
        table = self.doc.add_table(rows=rows, cols=cols)
        for r in range(rows):
            for c in range(cols):
                table.rows[r].cells[c].paragraphs[0].style = "No Spacing"
        return table

    def _create_results_table(self) -> Table:
        res = self._create_compact_table(1, 4)
        res.columns[0].width = Mm(80)  # 190
        res.columns[1].width = Mm(45)
        res.columns[2].width = Mm(40)
        res.columns[3].width = Mm(25)
        return res

    def _create_test_name_table(self) -> Table:
        res = self._create_compact_table(1, 2)
        res.columns[0].width = Mm(165)
        res.columns[1].width = Mm(25)
        return res

    def _generate_uberheader(self):
        ht = self._create_compact_table(1, 3)

        cell = ht.rows[0].cells[0].paragraphs[0].add_run()
        cell.add_picture('../frontend/img/logo_bw.png', width=Mm(50))


        # self.doc.add_picture('../frontend-reporter/public/img/logo.png', width=Mm(50))

        if self.szc.naglowek is not None:
            cell = ht.rows[0].cells[1].paragraphs[0]
            cell.add_run(self.szc.naglowek)
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.runs[-1].font.size = Pt(7)
            # add_float_picture(par, '../frontend-reporter/public/img/logo.png', width=Mm(50))

        cell = ht.rows[0].cells[2].paragraphs[0]
        barkod = (self.szc.dane_zlecenia['kodkreskowy'] or '').strip()
        if len(barkod) > 0 and len(barkod) <= 12:
            barkod_fn = f'/tmp/{barkod}.png'
            self.files_to_clean_up.append(barkod_fn)
            img = code128.image(barkod)
            img.save(barkod_fn)
            cell.add_run().add_picture(barkod_fn, width=Mm(40))
            cell.add_run("\n" + barkod)
        cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.runs[-1].font.size = Pt(7)

    def _generate_header(self):
        hdr = self.doc.add_heading('Laboratory test report', level=1)
        self.doc.add_heading(f"{self.szc.dane_zlecenia['imiona']} {self.szc.dane_zlecenia['nazwisko']}", level=2)

        table = self._create_compact_table(6, 2)

        def header(col, row, title, value):
            cell = table.rows[row].cells[col].paragraphs[0]
            cell.add_run(title + ': ')
            if not empty(value):
                value = str(value)
            else:
                value = 'no data'
            cell.add_run(value).bold = True
            for run in cell.runs:
                run.font.size = Pt(9)

        plec = self.szc.dane_zlecenia['plec']

        header(0, 0, 'Date of Birth', self.szc.dane_zlecenia['dataurodzenia'])
        header(0, 1, 'Address', (self.szc.dane_zlecenia['adres'] or '').replace('\r\n', '\n'))
        header(0, 2, 'Patient ID', self.szc.dane_zlecenia['pac_numer'])
        header(0, 3, 'Gender', {'K': 'F'}.get(plec, plec))
        header(0, 4, 'Order document ID', self.szc.dane_zlecenia['zewnetrznyidentyfikator'])
        header(0, 5, 'No./Date in the diagnostic work book',
               f"{self.szc.dane_zlecenia['numer']} / {self.szc.dane_zlecenia['datarejestracji']}")
        header(1, 0, 'Reffering Center', self.szc.dane_zlecenia['zleceniodawca'])
        header(1, 1, 'Ordering Physician', self.szc.dane_zlecenia['lek_nazwisko'])
        header(1, 2, 'Result is sent back to', 'Reffering Center')
        header(1, 3, 'Date and time of registration', self.szc.dane_zlecenia['godzinarejestracji'].strftime('%Y-%m-%d %H:%M'))
        header(1, 4, 'Date of the test', self.szc.data_badania)

    def _generate_sections(self):
        for sekcja in self.szc.tresc['Sekcje']:
            self._generate_section(sekcja)

    def _generate_section(self, sekcja):
        # print(sekcja['Styl'])
        self.sekcja = sekcja
        self.section_header_printed = False
        self.doc.add_heading(self.nazwa_sekcji(sekcja['Styl']), level=3)

        for probka in sekcja['Próbki']:
            self._generate_sample(probka)

    def _generate_section_header(self):
        if self.section_header_printed:
            return
        hdr = self._create_results_table()
        for i, t in enumerate(["Test name", "Test result", "Reference ranges", "Reference"]):
            run = hdr.rows[0].cells[i].paragraphs[0].add_run(t)
            run.font.size = Pt(7)
            if i == 1:
                hdr.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # run.font.color.rgb = RGBColor(0x3f, 0x2c, 0x36)
            # nie działa kolor
        self.section_header_printed = True

    def _generate_sample(self, probka):
        # print(probka['Materiał'], probka['Pobranie'], probka['Przyjęcie'])
        # Heading 9
        # Date and time collection: ……………… (collecting person…………)
        #
        # Date and time of material receipt: ………………….
        self.probka = probka
        if self.sekcja.get('Styl', {}).get('@drukowaćMateriały') == '1':
            self.section_header_printed = False
            par = self.doc.add_paragraph('Material: ', style="Body Text 3")
            nazwa_mat = self.nazwa_materialu(probka['Materiał'])
            par.add_run(nazwa_mat).font.bold = True
            if probka.get('Pobranie', {}).get('Godzina') is not None:
                par.add_run(', date and time of collection: ')
                par.add_run(probka.get('Pobranie', {}).get('Godzina').replace('T', ' ')).font.bold = True
            if probka.get('Pobranie', {}).get('Poborca') is not None:
                par.add_run(' (collecting person: ')
                par.add_run(
                    probka.get('Pobranie', {}).get('Poborca').get('Imiona', '') + ' ' + probka.get('Pobranie', {}).get(
                        'Poborca').get('Nazwisko', '')).font.bold = True
                par.add_run(')')
            if probka.get('Przyjęcie', {}).get('Godzina') is not None:
                par.add_run(', date and time of material receipt: ')
                par.add_run(probka.get('Przyjęcie', {}).get('Godzina').replace('T', ' ')).font.bold = True
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # print(self.sekcja.keys(), self.sekcja['Styl'])
            # print(probka.get('Materiał'))
            # print(probka.get('Pobranie'))
            # print(probka.get('Przyjęcie'))

        for wykonanie in probka['Wykonania']:
            self._generate_test_performance(wykonanie)

    def _generate_test_performance(self, wykonanie):
        self._generate_section_header()
        self.wykonanie = wykonanie
        bad = wykonanie['Badanie']
        self.bad_symbol = bad['Symbol']
        met = wykonanie['Metoda']
        self._proc = (met.get('Procedura') or {}).get('Nazwa')
        if bad['@drukowaćNazwę'] == 'zawsze':
            hdr = self._create_test_name_table()
            hdr.rows[0].cells[0].paragraphs[0].add_run(self.nazwa_badania(bad)).font.size = Pt(9)
            if self._proc is not None:
                hdr.rows[0].cells[1].paragraphs[0].add_run(self._proc).font.size = Pt(6)
                self._proc = None
        # print(wykonanie.keys())
        # print(wykonanie['Badanie'])
        for wynik in wykonanie['Wyniki']:
            self._generate_test_result(wynik)

    def _generate_reference_range(self, cell: BlockItemContainer, wyn: Dict[str, Any]):
        par = cell.paragraphs[0]
        norma = wyn.get('Norma', {})
        dn = norma.get('@drukowaćNormę')
        # print(norma)
        # ZakresOd, ZakresDo  - jak jest to i wynik liczbowy to można narysować
        if dn == 'opis':
            opis = norma.get('Opis')
            if opis:
                opis = opis.replace('\x97', '-')
                opis = self.tlumacz_norme(opis)
                par.add_run(opis).font.size = Pt(8)
        elif dn == 'tabelka':
            if norma.get('Tabelka'):
                tabelka = norma.get('Tabelka')
                tabelka = self.tlumacz_norme(tabelka)
                par.add_run(tabelka).font.size = Pt(8)
            # TODO właściwa tabelka, sprawdzić separator
        elif dn is None:
            pass
        else:
            raise RuntimeError(f"drukować normę: {dn}")
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _generate_test_result(self, wynik):
        # print(wynik)
        par = wynik['Parametr']
        # self.doc.add_paragraph('  ' + self.nazwa_parametru(par))
        table = self._create_results_table()
        result_text = self.tlumacz_wynik(wynik.get('WynikTekstowy'))
        if par['@drukowaćNazwę'] in ('gdyWynik',):
            cell = table.rows[0].cells[0].paragraphs[0]
            cell.add_run(self.nazwa_parametru(self.bad_symbol, par)).font.size = Pt(9)
        cell = table.rows[0].cells[1].paragraphs[0]
        cell.add_run(result_text).font.size = Pt(9)
        cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._generate_reference_range(table.rows[0].cells[2], wynik)
        cell = table.rows[0].cells[3].paragraphs[0]
        if self._proc is not None:
            cell.add_run(self._proc).font.size = Pt(6)
            self._proc = None
        else:
            cell.add_run()
        cell.runs[-1].font.size = Pt(6)

        # print(par)
        pass  # WynikTekstowy, Norma, Parametr

    def _generate_footer(self):
        pt = self._create_compact_table(2, 3)
        wykonanili = ['Performed by'] + self.szc.wykonali
        zatwierdzili = ['Approved by'] + self.szc.zatwierdzili
        cell = pt.rows[1].cells[0].paragraphs[0]
        cell.add_run("\n".join(wykonanili)).font.size = Pt(7)
        cell = pt.rows[1].cells[2].paragraphs[0]
        cell.add_run("\n".join(zatwierdzili)).font.size = Pt(7)

        if self.config.kliniczne:
            self.doc.add_paragraph(BADANIE_KLINICZNE)

        self.doc.add_paragraph(PODPIS)

        self.doc.add_paragraph(STOPKA, style="Body Text 3")

    def nazwa_sekcji(self, styl):
        if styl['Symbol'] in self._tlumaczenia['sek']:
            return self._tlumaczenia['sek'][styl['Symbol']]
        return styl['Nazwa']

    def nazwa_badania(self, bad):
        if bad['Symbol'] in self._tlumaczenia['bad']:
            return self._tlumaczenia['bad'][bad['Symbol']]
        n1 = bad['Nazwa']
        id = int(bad['@id'].split(':')[1])
        n2 = self.szc.nazwy.get(id, {}).get('bad_nazwaalternatywna')
        if not empty(n2) and n2.lower().strip() != n1.lower().strip():
            return f"{n2} / {n1}"
        return n1

    def nazwa_parametru(self, bad_symbol, par):
        spar = bad_symbol + '.' + par['Symbol']
        if spar in self._tlumaczenia['par']:
            return self._tlumaczenia['par'][spar]
        n1 = par['Nazwa']
        id = int(par['@id'].split(':')[1])
        n2 = self.szc.nazwy.get(id, {}).get('par_nazwaalternatywna')
        if not empty(n2) and n2.lower().strip() != n1.lower().strip():
            return f"{n2} / {n1}"
        return n1

    def nazwa_materialu(self, mat):
        if mat['Symbol'] in self._tlumaczenia['mat']:
            return self._tlumaczenia['mat'][mat['Symbol']]
        n1 = mat['Nazwa']
        id = int(mat['@id'].split(':')[1])
        n2 = self.szc.nazwy.get(id)
        if not empty(n2) and n2.lower().strip() != n1.lower().strip():
            return f"{n2} / {n1}"
        return n1

    def tlumacz_wynik(self, wyn):
        if str(wyn).upper() in self._tlumaczenia['wyn']:
            return self._tlumaczenia['wyn'][str(wyn).upper()]
        return wyn

    def tlumacz_norme(self, val):
        if val is None:
            return None
        for tr_from, tr_to in self._tlumaczenia['nor'].items():
            val = val.replace(tr_from, tr_to)
        return val

    def generate(self, fn):
        self.doc = Document()

        core_prop = self.doc.core_properties
        core_prop.title = 'Laboratory test report'
        core_prop.subject = 'original report: ' + self.szc.report_id
        core_prop.comments = 'generated by Alab Reporter / python-docx\nsource: ' + self.szc.report_id

        section = self.doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

        footer = section.footer  # a HeadersFooters collection object
        run = footer.paragraphs[0].add_run("Laboratory test report translation - page ")
        add_page_number(run)
        run.font.size = Pt(7)
        footer.paragraphs[0].aligment = WD_ALIGN_PARAGRAPH.RIGHT

        self._setup_styles()

        self._generate_uberheader()
        self._generate_header()
        self._generate_sections()
        self._generate_footer()

        # for style in self.doc.styles:
        #     if style.type == WD_STYLE_TYPE.PARAGRAPH:
        #         self.doc.add_paragraph(f"to jest styl {style.name}", style=style.name)
        #
        # p = self.doc.add_paragraph('A plain paragraph having some ')
        # p.add_run('bold').bold = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True
        #
        # self.doc.add_paragraph(
        #     'first item in unordered list', style='List Bullet'
        # )
        # self.doc.add_paragraph(
        #     'first item in ordered list', style='List Number'
        # )

        # self.doc.add_page_break()

        self.doc.save(fn)
        for fn in self.files_to_clean_up:
            if os.path.isfile(fn):
                os.unlink(fn)